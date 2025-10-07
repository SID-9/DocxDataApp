import re
from docx import Document
from .utils import post_process_candidate, next_non_empty_index

ENTITY_KEYWORDS = {
    "Counterparty": ["Counterparty"],
    "Initial Valuation Date": ["Initial Valuation Date"],
    "Notional": ["Notional", "Notional Amount", "Notional Amount (N)", "Amount"],
    "Valuation Date": ["Valuation Date"],
    "Maturity": ["Maturity", "Termination Date", "Expiry Date", "Tenor", "Duration"],
    "Underlying": ["Underlying", "Instrument", "Asset"],
    "Coupon": ["Coupon", "Coupon (C)", "Interest Rate"],
    "Barrier": ["Barrier", "Barrier (B)"],
    "Calendar": ["Calendar", "Business Day", "Schedule"]
}

ALL_KEYWORDS = sorted({k for kws in ENTITY_KEYWORDS.values() for k in kws}, key=len, reverse=True)


def extract_text_from_docx(docx_path):
    """Return list of lines: paragraphs and table rows converted to 'label\\tvalue' where possible."""
    doc = Document(docx_path)
    lines = []

    # paragraphs
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            for l in txt.splitlines():
                if l.strip():
                    lines.append(l.strip())

    #tables: process row-wise and pair adjacent cells where possible
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_cells:
                continue
            # if row has at least 2 cells, pair them as label\tvalue for adjacent pairs
            if len(row_cells) >= 2:
                # pair 0-1, 2-3, ...
                for i in range(0, len(row_cells), 2):
                    if i + 1 < len(row_cells):
                        lines.append(f"{row_cells[i]}\t{row_cells[i+1]}")
                    else:
                        lines.append(row_cells[i])
            else:
                # single cell row
                for l in row_cells:
                    lines.append(l)



    return lines

def extract_entities_from_lines(lines):
    extracted = {k: None for k in ENTITY_KEYWORDS.keys()}
    parties = {}

    # 1) Detect Party A / Party B preferentially from 'label\\tvalue' pairs
    for i, line in enumerate(lines):
        parts = [p.strip() for p in line.split('\t')]
        left = parts[0] if parts else line
        # if left is 'Party A' and right exists -> use it
        if re.match(r"^\s*Party\s*A\b", left, re.IGNORECASE):
            if len(parts) > 1 and parts[1].strip():
                parties["Party A"] = parts[1].strip()
            else:
                j = next_non_empty_index(lines, i)
                if j is not None:
                    parties["Party A"] = lines[j].strip()
        if re.match(r"^\s*Party\s*B\b", left, re.IGNORECASE):
            if len(parts) > 1 and parts[1].strip():
                parties["Party B"] = parts[1].strip()
            else:
                j = next_non_empty_index(lines, i)
                if j is not None:
                    parties["Party B"] = lines[j].strip()

    # Counterparty preference
    if "Party A" in parties:
        extracted["Counterparty"] = parties["Party A"]
    elif "Party B" in parties:
        extracted["Counterparty"] = parties["Party B"]
    # Inline "Counterparty: XYZ Bank" detection
    for i, line in enumerate(lines):
        if extracted["Counterparty"]:
            break
        m = re.search(r"\bCounterparty\s*[:\-]\s*(.+)", line, re.IGNORECASE)
        if m:
            extracted["Counterparty"] = m.group(1).strip()

    # 2) Extract other entities (use table pairs first, then inline, then label+next line)
    for i, line in enumerate(lines):
        parts = [p.strip() for p in line.split('\t')]
        for entity, keywords in ENTITY_KEYWORDS.items():
            if entity == "Counterparty":
                continue
            if extracted.get(entity) is not None:
                continue

            found = False
            # a) table-like: left matches keyword and right exists
            if len(parts) >= 2:
                left = parts[0]
                right = parts[1]
                for kw in keywords:
                    if re.match(rf"^\s*{re.escape(kw)}\b", left, re.IGNORECASE):
                        cand = post_process_candidate(right)
                        if cand:
                            extracted[entity] = cand
                            found = True
                            break
                if found:
                    continue

            # b) inline with explicit separators (colon/hyphen/tab) — do NOT use '(' as separator
            for kw in keywords:
                m_inline = re.search(rf"{re.escape(kw)}\s*(?:[:\-\–]|\t)\s*(.+)", line, re.IGNORECASE)
                if m_inline:
                    cand = post_process_candidate(m_inline.group(1))
                    if cand:
                        extracted[entity] = cand
                        found = True
                        break
            if found:
                continue

            # c) label-only line -> next non-empty line
            for kw in keywords:
                if re.match(rf"^\s*{re.escape(kw)}\s*(?:[:\-\–\(\t]|$)", line, re.IGNORECASE):
                    j = next_non_empty_index(lines, i)
                    if j is not None:
                        cand = post_process_candidate(lines[j])
                        if cand:
                            extracted[entity] = cand
                    break

    # Post-cleanup
    for k, v in list(extracted.items()):
        if isinstance(v, str):
            extracted[k] = v.replace("\t", " ").strip()

    # include parties for debug
    if parties:
        extracted["_Party A"] = parties.get("Party A")
        extracted["_Party B"] = parties.get("Party B")

    # added later
    if not any(extracted.values()):  # if everything is empty
        full_text = "\n".join(lines)
        patterns = {
            "Counterparty": r"(?:Counterparty|Party A)\s*(?:\(|:)?\s*([A-Z][A-Z\s&]+)",
            "Initial Valuation Date": r"Initial Valuation Date(?: of|:)?\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})",
            "Notional": r"Notional(?: amount)?(?: agreed)?(?: was|:)?\s*([A-Z]{3}\s*[0-9.,]+\s*(?:million|bn|billion|m|k)?)",
            "Valuation Date": r"Valuation Date(?: is|:)?\s*([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})",
            "Maturity": r"(?:matures on|Maturity(?: Date)?(?: is|:)?\s*)([0-9]{1,2}\s+[A-Za-z]+\s+[0-9]{4})",
            "Underlying": r"Underlying\s+(?:is|:)\s*(.+?)(?:\.\s|$)",
            "Coupon": r"Coupon(?: rate)?(?: is|:)?\s*([0-9.,]+%)",
            "Barrier": r"Barrier(?: is|:)?\s*([0-9.,]+%)",
            "Calendar": r"Calendar(?: is|:)?\s*([A-Z]+)"
        }
        for key, pattern in patterns.items():
            m = re.search(pattern, full_text, re.IGNORECASE)
            if m and not extracted.get(key):
                extracted[key] = m.group(1).strip()

    return extracted
